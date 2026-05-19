import io
import json
import os

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]

XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GOOGLE_SHEET_MIME = "application/vnd.google-apps.spreadsheet"
GOOGLE_DOC_MIME = "application/vnd.google-apps.document"
GOOGLE_SLIDE_MIME = "application/vnd.google-apps.presentation"


def get_drive_service():
    sa_json = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
    if sa_json:
        info = json.loads(sa_json)
    else:
        creds_file = os.getenv("GOOGLE_CREDENTIALS_FILE", "service-account.json")
        with open(creds_file) as f:
            info = json.load(f)
    creds = service_account.Credentials.from_service_account_info(info, scopes=SCOPES)
    return build("drive", "v3", credentials=creds)


def _build_parent_query(folder_ids: list[str]) -> str:
    """'id in parents' — OR'd for multiple folder IDs."""
    if not folder_ids:
        return ""
    parts = [f"'{fid}' in parents" for fid in folder_ids]
    return "(" + " or ".join(parts) + ")"


def list_files(folder_ids: list[str] | None = None, page_size: int = 100) -> list[dict]:
    """List files — optionally restricted to specific parent folder IDs."""
    if folder_ids is not None and len(folder_ids) == 0:
        return []

    service = get_drive_service()
    q_parts = ["trashed=false"]
    if folder_ids:
        q_parts.append(_build_parent_query(folder_ids))
    query = " and ".join(q_parts)

    results = (
        service.files()
        .list(
            pageSize=page_size,
            fields="files(id, name, mimeType, modifiedTime, size, parents)",
            orderBy="modifiedTime desc",
            q=query,
            includeItemsFromAllDrives=True,
            supportsAllDrives=True,
        )
        .execute()
    )
    return results.get("files", [])


def _download_bytes(service, file_id: str) -> bytes:
    """Download a binary (non-Google-native) file's raw bytes."""
    request = service.files().get_media(fileId=file_id)
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    return buffer.getvalue()


def _xlsx_to_text(xlsx_bytes: bytes) -> str:
    """Parse XLSX bytes into text — reads ALL sheets/tabs (not just the first)."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes), data_only=True, read_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"=== Sheet: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(cell.strip() for cell in cells):
                parts.append(",".join(cells))
        parts.append("")
    wb.close()
    return "\n".join(parts).strip()


def _docx_to_text(docx_bytes: bytes) -> str:
    """Parse DOCX bytes into plain text."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def read_file_content(file_id: str, mime_type: str) -> str:
    service = get_drive_service()

    # ── Spreadsheets — read ALL tabs ──
    if mime_type == GOOGLE_SHEET_MIME:
        # Native Google Sheet → export as XLSX (contains every tab),
        # then parse all sheets. CSV export only returns the first tab.
        data = service.files().export(fileId=file_id, mimeType=XLSX_MIME).execute()
        return _xlsx_to_text(data)

    if mime_type == XLSX_MIME:
        # Uploaded .xlsx → download binary → parse all tabs
        return _xlsx_to_text(_download_bytes(service, file_id))

    # ── Word documents ──
    if mime_type == DOCX_MIME:
        return _docx_to_text(_download_bytes(service, file_id))

    # ── Google Docs / Slides → plain-text export ──
    text_export = {
        GOOGLE_DOC_MIME: "text/plain",
        GOOGLE_SLIDE_MIME: "text/plain",
    }
    if mime_type in text_export:
        response = service.files().export(
            fileId=file_id, mimeType=text_export[mime_type]
        ).execute()
        return response.decode("utf-8") if isinstance(response, bytes) else response

    # ── Other files (txt, csv, md, ...) → raw download ──
    return _download_bytes(service, file_id).decode("utf-8", errors="ignore")


def search_files(query: str, folder_ids: list[str] | None = None, page_size: int = 40) -> list[dict]:
    if folder_ids is not None and len(folder_ids) == 0:
        return []

    service = get_drive_service()
    q_parts = [f"name contains '{query}'", "trashed=false"]
    if folder_ids:
        q_parts.append(_build_parent_query(folder_ids))
    q = " and ".join(q_parts)

    results = (
        service.files()
        .list(
            q=q,
            pageSize=page_size,
            fields="files(id, name, mimeType, modifiedTime, parents)",
            includeItemsFromAllDrives=True,
            supportsAllDrives=True,
        )
        .execute()
    )
    return results.get("files", [])


def file_in_folders(file_id: str, folder_ids: list[str]) -> bool:
    """Check a specific file has any of its parents in folder_ids."""
    if not folder_ids:
        return False
    service = get_drive_service()
    try:
        meta = (
            service.files()
            .get(fileId=file_id, fields="parents", supportsAllDrives=True)
            .execute()
        )
    except Exception:
        return False
    parents = set(meta.get("parents", []))
    return bool(parents.intersection(set(folder_ids)))
